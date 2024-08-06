import streamlit as st
import logging
from datetime import datetime
import pandas as pd
import json
import os
from io import StringIO, BytesIO
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns
from retrieval_response import retrieve_relevant_chunks, generate_response
from file_handler import extract_text_from_files
from processing import chunk_documents, vectorize_chunks, store_vectors_in_faiss

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Attempt to load the API key
try:
    api_key = st.secrets["gemini_api_key"]
    logger.info("API key loaded successfully.")
except KeyError:
    st.error("API key is missing. Please add it to the secrets.toml file.")
    api_key = None

# Define logo URL
logo_url = "https://i.pinimg.com/originals/5f/24/38/5f24384a518a90e30a2f1107141ab9d4.gif"  # Replace with your logo URL

def flatten(list_of_lists):
    """Flatten a list of lists into a single list."""
    return [item for sublist in list_of_lists for item in sublist]

def save_feedback_to_json(feedback_data, file_name='feedback.json'):
    if os.path.exists(file_name):
        with open(file_name, 'r') as f:
            existing_data = json.load(f)
    else:
        existing_data = []
    
    existing_data.extend(feedback_data)

    with open(file_name, 'w') as f:
        json.dump(existing_data, f, indent=4)

def export_to_text(chat_history):
    """Export chat history to a text file."""
    output = StringIO()
    for chat in chat_history:
        if chat['question']:
            output.write(f"User: {chat['question']}\n")
        output.write(f"Bot: {chat['answer']}\n\n")
    return output.getvalue()

def export_to_docx(chat_history):
    """Export chat history to a DOCX file."""
    doc = Document()
    for chat in chat_history:
        if chat['question']:
            doc.add_paragraph(f"User: {chat['question']}")
        doc.add_paragraph(f"Bot: {chat['answer']}\n")
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def load_feedback_from_json(file_name='feedback.json'):
    if os.path.exists(file_name):
        with open(file_name, 'r') as f:
            feedback_data = json.load(f)
        return feedback_data
    else:
        return []

def plot_feedback_analysis(feedback_data):
    df = pd.DataFrame(feedback_data)
    if not df.empty:
        df['feedback'] = df['feedback'].astype(str)
        feedback_counts = df['feedback'].value_counts()
        fig, ax = plt.subplots()
        sns.barplot(x=feedback_counts.index, y=feedback_counts.values, ax=ax)
        ax.set_title('Feedback Analysis')
        ax.set_xlabel('Feedback')
        ax.set_ylabel('Count')
        st.pyplot(fig)
    else:
        st.write("No feedback data available.")

def main():
    st.sidebar.markdown(f"""
        <div style="text-align: center;">
            <img src="{logo_url}" width="100" height="100">
        </div>
    """, unsafe_allow_html=True)
    
    st.sidebar.markdown("""
        <h1 style="font-size: 24px; text-align: center;">DocuChat</h1>
        <h2 style="font-size: 18px; text-align: center;">RAG-based Chatbot</h2>
    """, unsafe_allow_html=True)

    st.sidebar.title("Upload Files")
    uploaded_files = st.sidebar.file_uploader("Upload PDF/DOCX/TXT files", type=["pdf", "docx", "txt"], accept_multiple_files=True)

    # Initialize session state for chat history, feedback, and feedback file existence
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = [
            {"question": "", "answer": "Hi, I am your RAG-Based ChatBOT. Please upload the files if you didn't."}
        ]
    if 'feedback' not in st.session_state:
        st.session_state.feedback = []
    if 'feedback_file_exists' not in st.session_state:
        st.session_state.feedback_file_exists = False

    if uploaded_files:
        documents = []
        for uploaded_file in uploaded_files:
            st.sidebar.write(f"File uploaded: {uploaded_file.name}")
            documents.append(extract_text_from_files(uploaded_file))

        # Combine all extracted text from documents into one large text
        all_text = "\n\n".join(flatten(documents))

        # Process the combined text to create chunks
        chunks = chunk_documents(all_text.split('\n\n'))

        # Vectorize the chunks
        vectors, vectorizer = vectorize_chunks(chunks)

        # Store vectors in FAISS
        index = store_vectors_in_faiss(vectors)
        
        # Store the index, vectorizer, and chunks in session state for later use
        st.session_state.index = index
        st.session_state.vectorizer = vectorizer
        st.session_state.chunks = chunks

    user_query = st.chat_input("Ask a question:")
    if user_query:
        if 'index' in st.session_state and 'vectorizer' in st.session_state and 'chunks' in st.session_state:
            try:
                retrieved_chunks = retrieve_relevant_chunks(st.session_state.index, st.session_state.chunks, user_query, st.session_state.vectorizer)
                response = generate_response("\n\n".join(retrieved_chunks), user_query, api_key)
            except Exception as e:
                response = f"An error occurred while generating the response: {str(e)}"
        else:
            response = "Please upload files first."

        # Append the question and answer to the chat history
        st.session_state.chat_history.append({"question": user_query, "answer": response})

    # Display the chat history
    for i, chat in enumerate(st.session_state.chat_history):
        if chat['question']:
            with st.chat_message("user"):
                st.write(chat['question'])
        with st.chat_message("assistant"):
            st.write(chat['answer'])
            
            # Show feedback options only for valid questions
            if chat['question']:
                # Initialize the feedback radio button with no option selected by default
                feedback_options = ('Yes', 'No')
                
                feedback = st.radio(
                    "Was this response helpful?",
                    feedback_options,
                    key=f"feedback_{i}"  # Ensure unique key for each feedback radio button
                )

                # Record feedback only when a new query is submitted
                if st.button("Submit Feedback", key=f"submit_feedback_{i}"):
                    feedback_data = {
                        "question": chat['question'],
                        "answer": chat['answer'],
                        "feedback": feedback,
                        "timestamp": datetime.now().isoformat()
                    }
                    st.session_state.feedback.append(feedback_data)

                    # Show a success message after feedback submission
                    st.success("Your feedback has been sent!")

    # Save feedback to JSON file
    if st.session_state.feedback:
        save_feedback_to_json(st.session_state.feedback)
        st.session_state.feedback = []

    # Export chat history
    if st.sidebar.button("Export Chat History as Text"):
        chat_history_text = export_to_text(st.session_state.chat_history)
        st.sidebar.download_button("Download Chat History as Text", chat_history_text, "chat_history.txt", "text/plain", key='download-txt')
    
    if st.sidebar.button("Export Chat History as DOCX"):
        chat_history_docx = export_to_docx(st.session_state.chat_history)
        st.sidebar.download_button("Download Chat History as DOCX", chat_history_docx, "chat_history.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key='download-docx')
    
    # Feedback analysis button
    if st.sidebar.button("Show Feedback Analysis"):
        st.session_state.show_feedback_analysis = not st.session_state.get('show_feedback_analysis', False)
    
    if st.session_state.get('show_feedback_analysis', False):
        st.sidebar.subheader("Feedback Analysis")
        feedback_data = load_feedback_from_json()
        plot_feedback_analysis(feedback_data)

    # Watermark at the bottom
    st.markdown("<div style='text-align: center; font-size: 12px; color: gray;'>Developed by Rahul Mydur</div>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
